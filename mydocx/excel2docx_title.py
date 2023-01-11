from docx import Document
import openpyxl as xl
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt

base_path = '''C:\\Users\\heave\\Documents\\WeChat Files\\wxid_d6s6ieyj6kyd12\\FileStorage\\MsgAttach\\048c03ab555b4a3c30a28f70bcc859d4\\File\\2022-07\\'''


def read_excel():
    path = f'{base_path}采控平台建设内容.xlsx'
    ws = xl.load_workbook(path)
    sheets = ws.sheetnames
    wb = ws[sheets[0]]
    # merged_cells = wb.merged_cells
    # ordered_cells = get_merged_cells(merged_cells)
    # print_for(ordered_cells)
    lst = []
    for row in wb.rows:
        # print(row)
        row_lst = []
        for i in range(1, len(row)):
            cell = row[i]._value
            if cell is not None:
                row_lst.append(cell)
                # print(cell,end="\t")
            else:
                row_lst.append("---")
                # print("---",end="\t")
        lst.append(row_lst)
    lst.pop(0)
    return lst
    # print()
    # for cell in row:
    #     value = cell._value
    #     print(value)


def gen_docx_title():
    doc = Document()
    lst = read_excel()
    for row in lst:
        for i in range(len(row)):
            cell = row[i]
            if cell != '---':
                no = "1"
                for k in range(i):
                    no += ".1"
                head = doc.add_heading(level=i + 1)
                run = head.add_run(no + cell)
                run.font.color.rgb = RGBColor(0, 0, 0)
                # 设置英文和数字字体
                run.font.name = 'arial'
                # 设置中文字体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                run.font.size = Pt(16 - i * 1.5)
                run.bold = False  # 加粗
                run.italic = False  # 斜体
    doc.save('./test.mydocx')


def for_lst():
    lst = read_excel()
    for row in lst:
        for cell in row:
            # if cell != '---':

            print(cell, end="\t")
        print()


def get_merged_cells(lst):
    ordered_cells = []
    for item in lst:
        if not str(item).startswith("A"):
            ordered_cells.append(str(item))
    return sorted(ordered_cells)


def print_for(lst):
    for item in lst:
        print(item)


if __name__ == '__main__':
    kk = [3, 54, 763, 23, 4]
    dd = ['C26:C29', 'B2:B50', 'khg', 'B51:B74']
    # print(sorted(dd))
    # read_excel()
    # for_lst()
    gen_docx_title()
