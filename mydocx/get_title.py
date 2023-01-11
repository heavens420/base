import time

from docx import Document

# file_name = r'2-中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 传输网（CD级）-20220914.1'
# file_name = r'中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 IP城域、骨干_0915'
# file_name = r'7-中国电信新一代云网运营业务系统技术规范集采控中心系列功能、指令级接口技术要求 IPRAN STN --20220907 v1.0'
# file_name = r'1-中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 IP城域、骨干_0915'
# file_name = r'7-中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求IP新型城域网0922'
# file_name = r'7-中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求IP新型城域网0922'
file_name = r'11-中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 IPRAN STN --20220907 v1.0'

path = f'C:\\Users\\heave\\Desktop\\DeskTop\\考核相关专业api\\{file_name}' + '.docx'
# path = f'./download/{file_name}' + '.docx'
obj = Document(path)


def for_paragraphs():
    count = 0
    index = 0
    for content in obj.paragraphs:
        index += 1
        title = content.style.name
        name = content.text
        line_spacing = content.paragraph_format.line_spacing
        # if line_spacing == 1.5:
        #     print(f'{name}')
        # if name == '接口名称' and not str(title).startswith('Heading') and not str(title).endswith('标题'):
        #     print(content)
        # styles = obj.styles
        # for style in styles:
        #     print(style.name)
        # if str(title).endswith('标题') or str(title).startswith('Heading'):
        #     print(f'{name}---{title}')
        # if str(title).__contains__('接口名称'):
        #     print(f'{name}')

        # count = get_error_title(title, name, count, index)
        count = get_interface_name(title, name, index,count, obj.paragraphs)
        # count = get_all_title(title, name, count, index, obj.paragraphs)

        # print(title)
        # if name == '接口名称':
        #     print(f'{title}------------title---')
        # if title == '代码':
        #     print(f'{name}------------------')
        # if str(title).startswith('字母编号列项'):
        #     print(f'{name}------22222------------')
        # if title == '二级条标题':
        #     print(f'{name}------二级标题------------')
        # if title == '三级条标题':
        #     print(f'{name}------api说明------------')
        # if title == 'Normal (Web)':
        #     print(f'{name}------Method&Path------------')
        # if title == '正文表标题':
        #     print(f'{name}------表格上标题------------')
    print(count)


# 查找指定应是标题 不是标题的api
def get_interface_name(title, name, index, count, content):
    # api_name = '接口名称'
    api_name = '接口访问方法'
    if name == api_name and (not str(title).startswith('Heading') and not str(title).endswith('标题')):
        con = content[index - 1]
        line_spacing = con.paragraph_format.line_spacing
        count += 1
        value = content[index]
        print(f'{value.text}')
    return count


# 查找标题名称错误的api 并计数
def get_error_title(title, name, count, index):
    api_name = '接口名称'
    # api_name = '接口访问方法'
    if str(name).__contains__(api_name) and str(name).strip() != api_name and (
            str(title).startswith('Heading') or str(title).endswith('标题')):
        print(f'{obj.paragraphs[index].text}')
        count += 1
    return count


# 获取所有标题
def get_all_title(title, name, count, index, content):
    if name == '接口名称':
        count += 1
        value = content[index].text
        print(f'{value}')
    return count


if __name__ == '__main__':
    for_paragraphs()
