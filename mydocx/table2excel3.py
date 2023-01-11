from docx import Document
import openpyxl as xl
import os

# excel_path = r'C:\Users\420\Desktop\調研轉換結果.xlsx'
pre_path = "C:\\Users\\420\\Desktop\\kdgc\\new2\\"
suffix_path = "陕西反馈"
xlsx = ".xlsx"

excel_path = pre_path + suffix_path + xlsx

dir_path = r'C:\Users\420\Desktop\kdgc\new'

title_lst = ['省份', '类型', '地市', '设备名称', '设备厂家', '设备型号', '所在IDC及AS', '设备版本', '设备IP', 'LOOPBACK', '设备数量', '物理服务器还是虚拟主机',
             '物理设备型号配置/或虚拟机配置简况', '系统版本', '备注']


def get_all_file_names():
    file_lst1 = list()
    for root, dirs, files in os.walk(dir_path):
        # print(files, end='\n')
        # print(type(files))
        for fi in files:
            # print(fi[:-5])
            str1 = fi.split(".")[1]
            if str1 == "mydocx":
                file_lst1.append(os.path.join(root, fi))
        return file_lst1


file_lst = get_all_file_names()


# print(file_lst)


def get_title(obj):
    lst = list()
    for txt in obj.paragraphs:
        title = txt.style.name
        if str(title).startswith("Heading 3"):
            content = str(txt.text)
            if content.__contains__('IDC-') and not content.__contains__('）情况：'):
                lst.append(content)
    return lst


# lst = get_title()

def generate_title(ws):
    col = 1
    for item in title_lst:
        ws.cell(row=1, column=col, value=item)
        col += 1


# 填充行
def name2name(ws1, tit, con, row_num):
    for i in range(len(title_lst)):
        if title_lst[i] == tit:
            con = str(con).strip()
            ws1.cell(row=row_num, column=i + 1, value=con)


def device_type(type, ws, title2, content1, row_num, excel_name):
    if type.__contains__('设备列表'):
        # 路由器
        ws.cell(row=row_num, column=2, value='路由器')
    if type.__contains__('防火墙'):
        # 防火墙
        ws.cell(row=row_num, column=2, value='防火墙')
    if type.__contains__('交换机'):
        # 交换机
        ws.cell(row=row_num, column=2, value='交换机')
    if type.__contains__('质量测试服务器'):
        # 质量测试服务器
        ws.cell(row=row_num, column=2, value='质量测试服务器')
    if type.__contains__('其它'):
        # 其它
        ws.cell(row=row_num, column=2, value='其它')
    ws.cell(row=row_num, column=1, value=excel_name[-10:])
    name2name(ws, title2, content1, row_num)


def write_all_docx():
    for it in file_lst:
        excel_name = it[:-5]
        obj = Document(it)
        wb = xl.Workbook()
        tables = obj.tables
        foreach_all_tables(tables, wb, obj, excel_name)
        wb.save(excel_name + ".xlsx")


def foreach_all_tables(tables, wb, obj, excel_name):
    lst = get_title(obj)
    ws = wb.active
    generate_title(ws)
    row_num = 1
    for k in range(1, 7):
        table1 = tables[k]
        title_item = lst[k - 1]
        for i in range(1, len(table1.rows)):
            row_num = row_num + 1
            count = 3
            for j in range(0, len(table1.columns)):
                content1 = table1.cell(i, j).text
                title2 = table1.cell(0, j).text
                if count == len(table1.columns) - 1:
                    row_num -= 1
                if content1 is None or content1 == '无其他设备' or content1 == '无其他设备' == '无' or str(content1).strip() == '':
                    count += 1
                    continue
                device_type(title_item, ws, title2, content1, row_num, excel_name)
    print(f'总行数:{row_num}')


write_all_docx()
