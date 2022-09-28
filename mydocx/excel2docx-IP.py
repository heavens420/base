from docx import Document
import openpyxl as xl
from openpyxl.utils import get_column_letter

file_name = r'中国电信新一代云网运营业务系统技术规范集 采控中心接口系列 服务级接口技术要求 IP城域、骨干_0922'


# file_name = r'2-中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 传输网（CD级）-20220914.1'


# 创建docx对象
def get_docx_obj():
    path = f'C:\\Users\\heave\\Desktop\\DeskTop\\考核相关专业api\\{file_name}' + '.docx'
    doc = Document(path)
    return doc


# 创建openpyxl对象
def get_openpyxl_obj():
    wb = xl.Workbook()
    ws = wb.active
    return wb, ws


# 读取所有段落
def read_paragraphs():
    doc = get_docx_obj()
    for par in doc.paragraphs:
        print(par.text)


# 读取所有正文
def read_normal():
    doc = get_docx_obj()
    para = doc.paragraphs

    for pa in para:
        if str(pa.style.name).startswith('Normal'):
            print(pa.text)


# 读取所有标题
def read_title():
    doc = get_docx_obj()
    para = doc.paragraphs
    for content in para:
        title = content.style.name
        if str(title).startswith("Heading"):
            print(content.text)


# 读取所有表格
def read_all_tables():
    doc = get_docx_obj()
    tables = doc.tables
    count = 0
    for table in tables:
        row_len = len(table.rows)
        col_len = len(table.columns)
        result = judge_table(table)
        if result:
            count += 1
            for i in range(row_len):
                for j in range(col_len):
                    print(f'{table.cell(i, j).text}', end="\t")
                print()
            print(str(count) + '-' * 100 + str(count))


# 读取所有合规的表格
def read_tables() -> list:
    doc = get_docx_obj()
    tables = doc.tables
    # 三维数组
    tables_list = []
    for table in tables:
        result = judge_table(table)
        if result:
            # 表格
            table_list = []
            row_len = len(table.rows)
            col_len = len(table.columns)
            # for row in table.rows:
            for i in range(1, row_len):
                # 行
                row_list = []
                # for cell in table.cells:
                for j in range(col_len):
                    value = table.cell(i, j).text
                    # 剔除新旧值空列
                    if j == 3 and value == '':
                        continue
                    # 得到一行
                    row_list.append(value)
                    # 得到多行 即一个表格
                table_list.append(row_list)
                # 得到多个表格 即所有合规表格
            tables_list.append(table_list)
    print('获取表格结束')
    return tables_list


# 判断表头是否满足要求
def judge_table(table) -> bool:
    template = ['入参/出参', '参数位置','参数名称', '参数编码', '参数类型（string等）', '参数约束（必填m、可选o、条件可选c）', '参数说明（含字典值等）']
    template2 = ['出入参类型', '参数名称', '参数编码', '新旧值', '参数类型', '参数约束', '参数说明（含字典值等）']
    col_len = len(table.columns)
    for i in range(col_len):
        cell = str(table.cell(0, i).text).strip().lower()
        if cell in template or cell in template2:
            pass
        else:
            return False
    return True

    # compare = set(template) - set(first_row)
    # if len(compare) == 0:
    #     return True
    # return False


# 读取所有合规的文档信息(不包含表格)保存为列表
def read_doc() -> list:
    doc = get_docx_obj()
    para = iter(doc.paragraphs)
    doc_list = []
    head2 = '二级条标题'
    head3 = '三级条标题'
    head4 = '四级条标题'
    flag_interface_name = '接口名称'
    flag_method_name = ['接口访问方法']

    while 1:
        try:
            pa = next(para)
            row_list = []
            pa_type = str(pa.style.name).strip()
            name = str(pa.text).strip()
            # api功能简述
            if pa_type == head2:
                row_list = [name]
                # 走到下一个段 再判断是不是标题
                pa = next(para)
                pa_type = str(pa.style.name).strip()
                name = str(pa.text).strip()

                # 获取下一个段落 更新段落名称 直到是三级标题为止
                while not pa_type.endswith('标题'):
                    pa = next(para)
                    pa_type = str(pa.style.name).strip()
                    name = str(pa.text).strip()
                if name != flag_interface_name:
                    continue
                else:
                    nx = str(next(para).text)
                    interface_name = nx.split("：")[-1]
                    nx = str(next(para).text)
                    interface_code = nx.split("：")[-1]
                    row_list.append(interface_name)
                    row_list.append(interface_code)
                    # 获取下一个段落 更新段落名称
                    pa = next(para)
                    pa_type = str(pa.style.name).strip()
                    name = str(pa.text).strip()
                    while not pa_type.endswith('标题'):
                        pa = next(para)
                        pa_type = str(pa.style.name).strip()
                        name = str(pa.text).strip()
                # 判断是否是接口访问方法
                if name in flag_method_name:
                    method_name = ''
                    uri_name = ''
                    line = str(next(para).text).split(" ")
                    if len(line) == 2:
                        method_name = line[0]
                        uri_name = line[1]
                    elif len(line) == 1:
                        method_name = line[0]
                        uri_name = str(next(para).text).strip()
                    else:
                        method_name = line[0]
                        uri_name = line[-1]
                        print(f'接口访问方法格式异常：{line}')
                    row_list.append(method_name)
                    row_list.append(uri_name)
            # 五个参数俱全 满足返回要求
            if len(row_list) == 5:
                row_list[0], row_list[2] = row_list[2], row_list[0]
                doc_list.append(row_list)
        except StopIteration as e:
            print('获取文档内容(不包含表格)结束')
            return doc_list


# 第一行 标题行
def write_head_title(ws):
    head_title = ['RES API编码', 'RES API名称', 'API功能简述', 'METHOD', 'URI', '入参/出参', '参数名称', '参数编码', '参数类型（String等）',
                  '参数约束（必填M、可选O、条件可选C）', '参数说明（含字典值等）']
    width = 40
    for i in range(1, len(head_title) + 1):
        ws.cell(1, i, head_title[i - 1])

        # 调整列宽
        if i == 4:
            width = 10
        elif i == 5:
            width = 55
        elif i > 5:
            width = 20

        # 设置单列
        # ws.column_dimensions['A'].width = 20.0
        # 设置所有列
        ws.column_dimensions[get_column_letter(i)].width = width

    # 调整行高
    # 设置单行行高
    ws.row_dimensions[1].height = 30
    # 设置所有行行高
    # for i in range(1, ws.max_row + 1):
    #     ws.row_dimensions[i].height = height


# 合并文档内容 文字+表格
def merge_doc_table():
    wb, ws = get_openpyxl_obj()
    write_head_title(ws)

    doc_list = read_doc()
    # for kk in doc_list:
    #     print(f'{kk[2]}')
    tables_list = read_tables()
    if len(doc_list) != len(tables_list):
        print('文档合并失败：表格数和文档数不一致')

    # 行号
    row_number = 1
    for i in range(len(tables_list)):
        table = tables_list[i]
        write_doc = True
        for k in range(len(table)):
            row_number += 1
            if write_doc:
                # 1-5列
                for s in range(5):
                    ws.cell(row_number, s + 1, doc_list[i][s])
            # 6-11列
            for j in range(6, len(table[k])+6):
                write_doc = False
                ws.cell(row_number, j, table[k][j - 6])
        # write_doc = True

    wb.save(f'./download/{file_name}.xlsx')
    wb.close()


def test_iter():
    lst = [1, 2, 4]
    iterator = iter(lst)
    nx = next(iterator)

    while 1:
        try:
            if nx is None:
                break
            print(nx)
            nx = next(iterator)
        except StopIteration as e:
            return


if __name__ == '__main__':
    # read_paragraphs()
    # read_normal()
    # read_tables()
    # read_title()
    merge_doc_table()
    # test_iter()
    # ss = "2334"
    # ff = ss.split(":")
    # print(ff)
