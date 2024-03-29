from docx import Document
import openpyxl as xl
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches
import json
# from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT


# from mydocx.enum.table import wD_CELL_VERTICAL_ALIGNMENT

'''
    excel文档转word
'''
def read_excel():
    list_excel = []
    # excel_name = r'城域网三层VPN接入RFS_API参数.xlsx'
    target_path = r'C:\Users\heave\Documents\WeChat Files\wxid_d6s6ieyj6kyd12\FileStorage\MsgAttach\afe73556e8b34e0d695986ec49e1c667\File\2022-07\新增互联网专线+家宽故障诊断接口_0712.xlsx'
    wb = xl.load_workbook(target_path)

    ws = wb.active
    row_num, col_num = ws.max_row, ws.max_column

    for row in range(2, row_num + 1 + 1):
        row_list = []
        for col in range(1, col_num + 1):
            row_list.append(ws.cell(row, col).value)
        list_excel.append(row_list)
    # for kk in list_excel:
    #     for tt in kk:
    #         print(tt, end="\n")
    return list_excel


def handle_excel():
    global api
    dic = dict()
    param_list = []
    list_excel = read_excel()
    pre_api = 'first'
    # head_title = ['出入参类型', '参数名称', '参数编码', '新旧值', '参数类型', '参数约束', '默认值']
    # dic['head,hkdahdk'] = head_title
    for row in list_excel:
        if row[0] == pre_api or pre_api == 'first':
            params = []
            api = row[0] + ',' + row[1]
            for j in range(2, len(row)):
                if row[j] is None:
                    row[j] = ''
                params.append(row[j])
            param_list.append(params)
        else:
            param_list = []
            dic[api] = param_list
            # api = ''
            params = []
            if row[0] is not None:
                api = row[0] + ',' + row[1]
                for j in range(2, len(row)):
                    if row[j] is None:
                        row[j] = ''
                    params.append(row[j])
                param_list.append(params)

        pre_api = row[0]
    # for key in dic:
    #     print(str(key))
    #     for value in dic[key]:
    #         print(str(value),end="\t")
    #     print()
    return dic


def gen():
    doc = Document()
    handle_row(doc)

    # 全局样式
    # doc.styles["Normal"].font.name = u'宋体'
    # doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # doc.styles["Normal"].font.size = Pt(10.5)

    # doc.styles["Heading 4"].font.name = u'黑体'
    # doc.styles["Heading 5"].font.name = u'黑体'
    # doc.styles["Heading 4"].font.size = Pt(10.5)
    # doc.styles["Heading 5"].font.size = Pt(10.5)

    doc.save("./newFile.mydocx")


def gen_table(table_params, doc):
    # +1 是表头
    table = doc.add_table(rows=len(table_params) + 1, cols=6, style='Table Grid')
    # head_title = ['出入参类型', '参数名称', '参数编码', '新旧值', '参数类型', '参数约束', '默认值']
    head_title = ['入参/出参', '参数名称', '参数编码', '参数类型(String等)', '参数约束(必填M,可选O,条件可选C)', '参数说明(含字典值等)']
    # 设置表头
    # cells = table.rows[0].cells
    for j in range(6):
        # cells[j].text = str(head_title[j])
        # table.cell(0, j).text = str(head_title[j])
        run = table.cell(0, j).paragraphs[0].add_run(head_title[j])
        run.font.name = 'arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    # 填充表数据 数据从第二行开始填充 防止覆盖表头
    for i in range(len(table_params)):
        cells = table.rows[i + 1].cells
        for j in range(6):
            # cells[j].text = str(table_params[i][j])
            run = cells[j].paragraphs[0].add_run(str(table_params[i][j]))
            run.font.name = 'arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    table.style.font.size = Pt(10.5)
    # table.style.font.color.rgb = RGBColor(255, 0, 0)
    table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def handle_row(doc):
    dic = handle_excel()
    # 五级标题
    no = 0
    for key in dic:
        no += 1
        # 行下五级标题
        # 标题前的数字编号标题
        no_title = f"6.3.1.{no}"
        # 标题下级标题 接口名称
        api_title = key.split(",")[0]
        # 标题
        title = str(api_title).replace("_api", "接口")

        # 接口名称 五级标题
        name_title = '接口名称'

        # 接口名称同级标题 接口编码
        api_code = key.split(",")[1]

        # 接口访问方法
        method_title = '接口访问方法'

        method = r'    POST'
        new_api_code = str(api_code).split("_")
        prefix = new_api_code[0]
        mid = new_api_code[1]
        tail = new_api_code[2]
        tail2 = new_api_code[3]
        tail3 = new_api_code[4]
        url = ''
        if len(new_api_code) == 6:
            tail4 = new_api_code[5]
            url = f'    /api/rest/ctrl/ip/{tail}/{tail2}/{tail3}/{tail4}'.lower()
        if len(new_api_code) == 5:
            url = f'    /api/rest/ctrl/ip/{tail}/{tail2}/{tail3}'.lower()
        # url = f'    /api/rest/ctrl/ip/{prefix}/{mid}/{tail}'.lower()

        # api 定义方法 字体红色
        define_title = 'API定义方法'
        define_content = '''    补充内容：入参、出参（必须定义出错时的返回值和说明）、报文样例等 可以定义通用接口出错返回值，有特殊的要单独定义'''

        # 设置样式
        # 样式 四级标题 eg:6.3.1.1　CN2PE设备RES_城域网-CN2三层VPN接入RFS_装接口
        set_head_style(doc, no_title + title, 4, 14)

        # eg: 接口名称
        set_head_style(doc, no_title + ".1" + name_title, 5, 12)

        # api名称 api编码内容
        set_content(doc, "    API名称：" + api_title + "\n    API编码：" + api_code, par_indent=True)

        # eg: 接口访问方法
        set_head_style(doc, no_title + ".2" + method_title, 5, 12)

        set_content(doc, method + "\n" + url, par_indent=True)
        # set_content(doc, url)

        # API定义方法
        set_head_style(doc, no_title + ".3" + define_title, 5, 12)

        set_content(doc, define_content, first_indent=True)
        # doc.add_paragraph(define_content).paragraph_format.left_indent = Cm(0.75)
        # 表格
        table_title = f'表1　{api_title}参数表'
        set_content(doc, table_title, False, True)
        # doc.add_paragraph(table_title).paragraph_format.left_indent = Cm(0.75)

        # 表格参数数据
        table_params = dic[key]
        # 画表格
        gen_table(table_params, doc)

        req_title = "请求示例"
        set_head_style(doc, no_title + ".4" + req_title, 5, 12)
        a = "    a)发送报文样例:"
        # URI = method + "  " + "http://xxx:9999/" + url
        URI = method + "   " + "http://xxx:9999" + url.replace(" ", "")
        accept = "    Accept: application/yang-data+json"
        params_json = gen_json(table_params, api_title, '入参')
        out_params_json = gen_json(table_params, api_title, '出参')
        #     b = """
        # b )应答示例：
        #     HTTP/1.1 201 Created
        #     Date: Thu, 26 Jan 2022 20:56:30 GMT
        #     Server: example-server
        #     Cache-Control: no-cache
        # e )错误码说明："""
        b = f"""
    b )应答示例：
        {out_params_json}
    e )错误码说明："""
        biao2 = "Xx错误编码"

        set_content(doc, a + "\n" + URI + "\n" + accept)
        # set_content(doc, URI)
        # set_content(doc, accept)
        set_content(doc, params_json)
        set_content(doc, b)
        set_content(doc, biao2, center=True)

        # 错误编码表格
        biao_list = [
            ['错误编码', '说明'],
            ['IpInBlacklist', 'IP地址在黑名单中']
        ]
        table2 = doc.add_table(rows=len(biao_list), cols=2, style='Table Grid')

        for i in range(len(biao_list)):
            for j in range(2):
                # table2.cell(i, j).text = biao_list[i][j]
                run = table2.cell(i, j).paragraphs[0].add_run(biao_list[i][j])
                run.font.name = 'Arial'  # 英文字体设置
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


# 整理一个接口函数，可以扩充完善
def set_head_style(doc, head_content, level, size):
    """
    doc: word文档对象
    head_content: 标题的内容
    level: 标题的等级，一级标题、二级标题、三级标题
    size: 标题的大小
    """
    head = doc.add_heading("", level=level)
    run = head.add_run(head_content)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.bold = False  # 加粗
    run.italic = False  # 斜体


def set_content(doc, content, color=False, center=False, first_indent=False, par_indent=False):
    pa = doc.add_paragraph()
    run = pa.add_run(content.replace("  ", ""))
    # 首行缩进
    if first_indent:
        pa.paragraph_format.first_line_indent = Inches(0.3)
    # 段落缩进
    if par_indent:
        pa.paragraph_format.left_indent = Inches(0.3)
    run.font.name = 'arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 设置样式
    # doc.add_paragraph().paragraph_format.left_indent = Cm(0.75)  # 段首缩进（word只有左侧才将段首缩进）
    # run.font.name = u'宋体'
    run.font.size = Pt(10.5)
    # if color:
    #     run.font.color.rgb = RGBColor(255, 0, 0)
    # 段落居中如何设置
    if center:
        pa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# 二维数组转json
def gen_json(table_params, api_name, param_type):
    index = 0
    # 最后要补充的右括弧数量
    kuohu_count = 0
    # 初始化的左括弧
    param_json = "{"
    # 上一个参数级别
    pre_level = 0
    # 遍历二维表
    for params in table_params:
        index += 1
        # 获取参数级别
        level = param_level(params[2])
        # 格式化参数
        param = str(params[2]).replace(">", "")
        # 只转换入参 出参同理
        if str(params[0]).strip() == param_type:
            # 如果当前参数级别大于上级 即当前参数可能为上级参数的子参数 即嵌套关系 （标题越小 级别越高层级越靠外 原理类似word标题级别）
            if level >= pre_level:
                # 构建json 往后加参数
                param_json += f"\"{param}\""
                # 判断当前参数是不是上级参数子参数
                if str(params[3]).strip().lower() == "Object".lower() or \
                        str(params[3]).strip().lower() == "Array".lower():
                    # 判断当前参数是否有子参数 有 加 {
                    param_json += ":{"
                    # 同时右括弧数量 +1
                    kuohu_count += 1
                else:  # 当前参数没有子参数 直接在参数后面补上 : "",
                    # param_json += ":\"\","
                    param_json += ":\""+str(params[5])+"\","
            else:  # 当前参数级别小于上个参数 说明 上一个json对象已经结束 下面开始下一级json对象的构造
                # 所以要给上级json (子)对象补上 右括弧
                for i in range(pre_level - level):
                    param_json = param_json + "},"
                    # 同时 总的右括弧数量要 -1
                    kuohu_count -= 1
                # 这里就是下一个json对象的构建了 原理同上
                param_json += f"\"{param}\""
                if str(params[3]).strip().lower() == "Object".lower() or \
                        str(params[3]).strip().lower() == "Array".lower():
                    param_json += ":{"
                    kuohu_count += 1
                else:
                    # param_json += ":\"\","
                    param_json += ":\""+str(params[5])+"\","

        # 把当前参数等级保存 用于上级参数和当前参数的比较
        pre_level = level
        # 无用代码 根本不会执行 但可优化代码 使之有用
        if pre_level == 0:
            kuohu_count = 0
    # 补全最后的右括弧
    for i in range(kuohu_count):
        param_json = param_json + "},"
    # 对应初始化的 {
    param_json = param_json + "}"
    # 处理所有对象的逗号问题
    param_json = param_json.replace(",}", "}")
    # 格式化json对象
    try:
        param_json = json.loads(param_json)
        param_json = json.dumps(param_json, indent=4, ensure_ascii=False)
    except:
        print('这里打印的是格式化出错的json')
        print(f'index:{index} {api_name}')
        print(param_json)

    return param_json


def test_json():
    dic = handle_excel()
    for key in dic:
        value = dic[key]
        format_json = gen_json(value)
        with open("./json_test.json", "w+") as jn:
            jn.write(format_json)
        break


# 判断参数的层级
def param_level(param):
    level = 1
    for i in range(len(param)):
        if param[i] == ">":
            level += 1
        else:
            break
    return level


if __name__ == '__main__':
    gen()
    # read_excel()
    # handle_excel()
    # test_json()
