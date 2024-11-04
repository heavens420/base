import random
import time

from docx import Document
import openpyxl as xl
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl.utils import get_column_letter

# file_name = r'12-中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 PON 0922(1)'
# file_name = r'10-中国电信新一代云网运营业务系统技术规范集 采控中心系列 服务级接口技术要求 IPRAN STN- -20220907 v1.0'
# file_name = r'1-中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 IP城域、骨干_0915 - 副本 (2)'
# file_name = r'中国电信新一代云网运营业务系统技术规范集-采控中心系列-服务级接口技术要求IP新型城域网-0909'
# file_name = r'7-中国电信新一代云网运营业务系统技术规范集采控中心系列功能、指令级接口技术要求 IPRAN STN --20220907 v1.0'
# file_name = r'中国电信新一代云网运营业务系统技术规范集采控中心系列服务级接口技术要求 IPRAN STN- -20220907 v1.0'
# file_name = r'中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 IP城域、骨干_0915'
# file_name = r'中国电信新一代云网运营业务系统技术规范集-采控中心接口系列-服务级接口技术要求 IP城域、骨干_0915'
# file_name = r'2-中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 传输网（CD级）-20220914.1'

# file_name = r'中国电信新一代云网运营业务系统采控中心接口服务级接口技术要求 IP城域、骨干_20230418'
# file_name = r'8中国电信新一代云网运营业务系统技术规范集 采控中心系列 功能、指令级接口技术要求 IP城域、骨干'
# file_name = r'中国电信新一代云网运营业务系统采控中心接口服务级接口技术要求 IP城域、骨干_20230418'
# file_name = r'中国电信新一代云网运营业务系统采控中心服务级接口技术要求 IP城域、骨干_20230823'
# file_name = r'中国电信新一代云网运营业务系统采控中心功能、指令级接口技术要求 IP城域、骨干_20230912（含附录C）'
# file_name = r'中国电信新一代云网运营业务系统技术规范集 采控中心系列 服务级接口技术要求 云0810 -自研虚拟化(1)'
# file_name = r'中国电信新一代云网运营业务系统采控中心技术要求 5GC 服务级接口 -20221222'
# file_name = r'中国电信新一代云网运营业务系统采控中心技术要求 5GC 功能、指令级接口-20220525-20230625 - 20230910(1)'
# file_name = r'中国电信新一代云网运营业务系统采控中心技术要求 5GC 功能、指令级接口-20231215'
# file_name = r'2-中国电信新一代云网运营业务系统采控中心功能、指令级接口技术要求 传输网（CD级）-20221017.1版本-20221026更新 - 20221214格式调整'
# file_name = r'3-中国电信新一代云网运营业务系统采控中心功能、指令级接口技术要求 无线网-1013'
# file_name = r'4-中国电信新一代云网运营业务系统采控中心技术要求 5GC 服务级接口 -20221222'
# file_name = r'5-中国电信新一代云网运营业务系统采控中心技术要求 5GC 功能、指令级接口-20221219'

# file_name = r'2024-01-11中国电信新一代云网运营业务系统采控中心技术要求-IP新城 服务级接口'
file_name = r'2024-01-11中国电信新一代云网运营业务系统采控中心技术要求-IP新城 功能、指令级接口'
# path = f'C:\\Users\\heave\\Desktop\\DeskTop\\考核相关专业api\\{file_name}' + '.docx'
path = f'C:\\Users\\heave\\Desktop\\mysource\\0919规范收集-word转excel\\{file_name}' + '.docx'
# path = f'C:\\Users\\heave\\Desktop\\{file_name}' + '.docx'
'''
    pip install python-docx
    pip install openpyxl
'''


# 创建docx对象
def get_docx_obj():
    doc = Document(path)
    return doc


# 创建openpyxl对象
def get_openpyxl_obj():
    wb = xl.Workbook()
    ws = wb.active
    return wb, ws


# 读取所有段落
def read_paragraphs():
    doc = get_docx_obj()
    for par in doc.paragraphs:
        print(par.text)


# 读取所有正文
def read_normal():
    doc = get_docx_obj()
    para = doc.paragraphs

    for pa in para:
        if str(pa.style.name).startswith('Normal'):
            print(pa.text)


# 读取所有标题
def read_title():
    doc = get_docx_obj()
    para = doc.paragraphs
    for content in para:
        title = content.style.name
        if str(title).startswith("Heading"):
            print(content.text)


# 读取所有表格
def read_all_tables():
    doc = get_docx_obj()
    tables = doc.tables
    count = 0
    for table in tables:
        row_len = len(table.rows)
        col_len = len(table.columns)
        result = judge_table(table)
        if result:
            count += 1
            for i in range(row_len):
                for j in range(col_len):
                    print(f'{table.cell(i, j).text}', end="\t")
                print()
            print(str(count) + '-' * 100 + str(count))


# 读取所有合规的表格
def read_tables() -> list:
    doc = get_docx_obj()
    tables = doc.tables

    begin = time.time()
    total_time, judge_total_time = 0, 0

    # 三维数组
    tables_list = []
    for table in tables:
        judge_time = time.time()
        result = judge_table(table)
        judge_end_time = time.time()
        judge_total_time += (judge_end_time - judge_time)
        if result:
            # 表格
            table_list = []
            row_len = len(table.rows)
            col_len = len(table.columns)
            for i in range(1, row_len):
                # 行
                row_list = []
                for j in range(col_len):
                    value = table.cell(i, j).text
                    # 剔除新旧值空列
                    # if j == 3 and value == '':
                    #     continue
                    # 得到一行
                    row_list.append(value)
                    # 得到多行 即一个表格
                if len(row_list) < 6:
                    continue
                table_list.append(row_list)
                # 得到多个表格 即所有合规表格
            tables_list.append(table_list)
    print('获取表格结束')
    end_time = time.time()
    total_time = end_time - begin - judge_total_time
    print(f'total_time:{total_time / 1000}')
    print(f'judge_time:{judge_total_time / 1000}')
    return tables_list


# 判断表头是否满足要求
def judge_table(table) -> bool:
    template = ['入参/出参','回参/出参','入参/回参','回参/回参','回参', '参数位置', '参数名称', '参数编码', '参数类型（string等）', '参数类型(string等)', '参数类型(string等）',
                '参数约束（必填m、可选o、条件可选c）', '参数约束(必填m,可选o,条件可选c)', '参数约束（必m、可选o、条件可选c）', '参数说明（含字典值等）']
    template2 = ['出入参类型', '参数名称', '参数编码', '新旧值', '参数类型', '参数约束', '参数说明(含字典值等)', '默认值', '参数示例', '备注', '']
    col_len = len(table.columns)
    if col_len < 6:
        return False
    # col_len = 2
    for i in range(col_len):
        cell = str(table.cell(0, i).text).strip().lower().replace('\n', '').replace('\t', '').replace('\u3000', '')
        # if len(table.columns) < 6:
        #     print(f'不合规表头具体项：{cell},列宽：{col_len}')
        #     print('*' * 100)
        #     for_table(table)
        #     print('-' * 100)
        if cell in template or cell in template2:
            if len(table.columns) < 6:
                print(f'不合规表头具体项：{cell},列宽：{col_len}')
                for_table(table)
                return False

            #     print('#' * 100)
            #     for_table(table)
            #     print('-1' * 100)
            #     print('-2' * 100)
            #     print('-3' * 100)
            #     print('-4' * 100)
            elif table.cell(0, 0).text == table.cell(0, 1).text == table.cell(0, 2).text == '':
                return False
            else:
                return True
        else:
            if len(table.columns) >= 6:
                print(f'不合规表头具体项：{cell},列宽：{col_len}')
                print('*' * 100)
                for_table(table)
                print('-' * 100)
            return False
    return True

    # compare = set(template) - set(first_row)
    # if len(compare) == 0:
    #     return True
    # return False


# 遍历table
def for_table(table):
    for row in table.rows:
        for cell in row.cells:
            print(f'{cell.text}', end='\t')
        print()


def for_list(lst):
    for it in lst:
        print(it)


# 读取所有合规的文档信息(不包含表格)保存为列表
def read_doc() -> list:
    doc = get_docx_obj()
    para = iter(doc.paragraphs)
    doc_list = []
    flag_interface_name = '接口名称'
    flag_method_name = ['接口访问方法']
    pre_name = ''

    while 1:
        try:
            pa = next(para)
            row_list = []
            pa_type = str(pa.style.name).strip()
            name = str(pa.text).strip()
            line_spacing = pa.paragraph_format.line_spacing

            if name == 'BGP-LS配置查询':
                pass
            # api功能简述
            if pa_type.endswith('标题') or pa_type.startswith('Heading') or line_spacing == 1.5 or pa_type.startswith(
                    'List'):
                if name != flag_interface_name:
                    pre_name = name
                    continue
                else:
                    row_list = [pre_name]
                    nx = str(next(para).text)
                    interface_name = nx.split("：")[-1]
                    nx = next(para)
                    interface_code = str(nx.text).split("：")[-1]
                    name = str(nx.text).strip()
                    pa_type = str(nx.style.name).strip()
                    line_spacing = nx.paragraph_format.line_spacing
                    row_list.append(interface_name)

                    if not (pa_type.endswith('标题') or pa_type.startswith('Heading') or line_spacing == 1.5):
                        row_list.append(interface_code)
                        # 获取下一个段落 更新段落名称
                        pa = next(para)
                        pa_type = str(pa.style.name).strip()
                        name = str(pa.text).strip()
                        line_spacing = pa.paragraph_format.line_spacing

                        while not (pa_type.endswith('标题') or pa_type.startswith('Heading') or line_spacing == 1.5):
                            pa = next(para)
                            pa_type = str(pa.style.name).strip()
                            name = str(pa.text).strip()
                            line_spacing = pa.paragraph_format.line_spacing
                            if name == 'API定义方法' or name == 'API参数定义':
                                break
                    else:
                        row_list.append(interface_name)
                # 判断是否是接口访问方法
                if name in flag_method_name:
                    method_name = ''
                    uri_name = ''
                    nx = next(para)
                    if str(nx.text).__contains__('\t'):
                        line = str(nx.text).split('\t')
                    elif str(nx.text).__contains__(':'):
                        line = str(nx.text).split(':')
                    elif str(nx.text).__contains__('：'):
                        line = str(nx.text).split('：')
                    else:
                        line = str(nx.text).split(" ")
                    if len(line) == 2:
                        method_name = line[0]
                        uri_name = line[1]
                    elif len(line) == 1:
                        method_name = line[0]
                        uri_name = str(next(para).text).strip()
                    else:
                        while '' in line and len(line) > 0:
                            line.remove('')
                        method_name = line[0]
                        uri_name = line[-1]
                        if len(line) != 2:
                            print(f'接口访问方法格式异常：{nx}')
                    row_list.append(method_name)
                    row_list.append(uri_name)
            # 五个参数俱全 满足返回要求
            if len(row_list) >= 5:
                row_list[0], row_list[2] = row_list[2], row_list[0]
                doc_list.append(row_list)
        except StopIteration as e:
            print('获取文档内容(不包含表格)结束')
            return doc_list


# 第一行 标题行
def write_head_title(ws):
    head_title = ['RES API编码', 'RES API名称', 'API功能简述', 'METHOD', 'URI', '入参/出参', '参数名称', '参数编码', '参数类型（String等）',
                  '参数约束（必填M、可选O、条件可选C）', '参数说明（含字典值等）']
    width = 40
    for i in range(1, len(head_title) + 1):
        ws.cell(1, i, head_title[i - 1])

        # 调整列宽
        if i == 4:
            width = 10
        elif i == 5:
            width = 55
        elif i > 5:
            width = 20

        # 设置单列
        # ws.column_dimensions['A'].width = 20.0
        # 设置所有列
        ws.column_dimensions[get_column_letter(i)].width = width

    # 调整行高
    # 设置单行行高
    ws.row_dimensions[1].height = 30
    # 设置所有行行高
    # for i in range(1, ws.max_row + 1):
    #     ws.row_dimensions[i].height = height


# 合并文档内容 文字+表格
def merge_doc_table():
    wb, ws = get_openpyxl_obj()
    write_head_title(ws)
    count = 0

    doc_list = read_doc()
    # for kk in doc_list:
    #     print(f'{kk[1]}')
    tables_list = read_tables()
    length = max(len(doc_list), len(tables_list))
    if len(doc_list) != len(tables_list):
        print('文档合并失败：表格数和文档数不一致')

    # 行号
    row_number = 1
    for i in range(length):
        table = []
        try:
            table = tables_list[i]
            write_doc = True

            # 没有表格参数 但是doc还是得写
            if len(table) == 0:
                row_number += 1
                for s in range(5):
                    try:
                        ws.cell(row_number, s + 1, doc_list[i][s])
                    except Exception as e:
                        print(f'{e},{doc_list[i]},i={i},s={s}')

            # 根据表格参数写 表格数据和doc数据
            for k in range(len(table)):
                row_number += 1
                if write_doc:
                    # 1-5列
                    for s in range(5):
                        try:
                            ws.cell(row_number, s + 1, doc_list[i][s])
                        except Exception as e:
                            print(f'{e},{doc_list[i]},i={i},s={s}')

                # 6-11列
                for j in range(6, len(table[k]) + 6):
                    try:
                        write_doc = False
                        ws.cell(row_number, j, table[k][j - 6])
                    except Exception as e:
                        print(f'i = {i},j = {j} 表格数据异常' + '-' * 30, e)
                        for_list(table)
        except Exception as e:
            count += 1
            # print('表格数据异常' + '-' * 30, e)
            # for_list(table)

    wb.save(f'./download/{file_name}.xlsx')
    wb.close()
    print(f'表格数据异常个数：{count}')


def test_iter():
    lst = [1, 2, 4]
    iterator = iter(lst)
    nx = next(iterator)

    while 1:
        try:
            if nx is None:
                break
            print(nx)
            nx = next(iterator)
        except StopIteration as e:
            return


# 格式化标题
def format_title(doc, content, level='5'):
    # head = doc.add_heading("", level=5)
    # run = head.add_run(content)
    # run.font.size = Pt(10.5)
    # run.font.color.rgb = RGBColor(0, 0, 0)
    # run.font.name = 'arial'
    # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    # run.bold = False  # 加粗
    # run.italic = False  # 斜体
    content.style.name = 'Heading ' + level
    print(level)
    # 修改样式
    content.paragraph_format.line_spacing = 1.5
    for run in content.runs:
        run.font.size = Pt(10.5)
        run.font.name = u'黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        run.font.italic = False
        run.font.bold = False


# 修改文档所有不合规标题
# 暂不可用
def for_paragraphs():
    doc = get_docx_obj()
    line = doc.paragraphs

    for content in line:
        title = content.style.name
        name = str(content.text)
        format_title(doc, content, str(random.randint(1, 5)))
        if name == '接口名称' and (not str(title).startswith('Heading') and not str(title).endswith('标题')):
            # doc1.add_heading(name, level=6)
            format_title(doc, content)
            # ss = dir(content.paragraph_format)
            # content.paragraph_format.style.name = 'Heading 5'
            # temp_content.style.name = 'Heading 5'
        if name == '接口访问方法' and (not str(title).startswith('Heading') and not str(title).endswith('标题')):
            # doc1.add_heading(name, level=6)
            format_title(doc, content, '6')
            # temp_content.style.name = 'Heading 5'
            # content.paragraph_format.style.name = 'Heading 5'
            # pass
    doc.save(f'./download/{file_name}-new.docx')


def get_all_number():
    doc = Document(path)
    for p in doc.paragraphs:
        # 获取numId
        # print('numId', p._element.pPr.numPr.numId.val, end='  ')
        # print('numId', p._element.pPr.numPr, end='  ')
        # num_pr = p._p.pPr.numPr
        # print(num_pr)
        if p.text == '认证及获取令牌':
            st = p.style
            print(st)

        if p.style.name.startswith('List'):
            # if p is not None and p.run is not None:
            #     print(f'no: {p.run[0].text}')
            # else:
            #     pass
            # print(f'no2: {p.runs}')

            # 获取ilvl的值，注意纯文本段落没有ilvl，其ilvl是None
            # try:
            #     print('ilvl', p._element.pPr.numPr.ilvl.val, end='  ')
            # except AttributeError:
            #     print('ilvl', p._element.pPr.numPr.ilvl, end='  ')
            # 获取每个段落的文本信息
            print('text', p.text)


if __name__ == '__main__':
    # read_paragraphs()
    # read_normal()
    # lst = read_tables()
    # for t in lst:
    #     print(t)
    # read_title()
    merge_doc_table()
    # for_paragraphs()

    # get_all_number()
